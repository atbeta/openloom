"""Format-aware task parsing and result/status rendering.

Supported input formats: JSON, YAML, DOCX (2-column table).
Output format always matches input format for both status and result files.
"""

from __future__ import annotations

import io
import json
import logging
import os
from datetime import UTC, datetime
from pathlib import PurePosixPath
from typing import Any
from zoneinfo import ZoneInfo

TASK_EXTENSIONS = frozenset({".json", ".yaml", ".yml", ".docx"})


def _resolve_tz() -> ZoneInfo | None:
    """Resolve the display timezone for docx timestamps.

    Order: env ``OPENLOOM_TIMEZONE`` > config file's ``harness.timezone``
    > system local. Default to system local — openloom usually runs on
    the operator's own machine, so the docx is for the operator, and
    the operator's wall-clock is what they expect to see. Override
    only when needed.

    Returns ``None`` when the system local timezone should be used;
    in that case ``datetime.fromtimestamp(ts, tz=None)`` falls through
    to the platform's ``localtime`` (per CPython docs since 3.6).
    """
    env = os.getenv("OPENLOOM_TIMEZONE", "").strip()
    if env:
        try:
            return ZoneInfo(env)
        except Exception:
            pass  # bad name → fall through to config / default
    try:
        from openloom.core.settings_source import find_config_file
        import yaml
        cfg_path = find_config_file()
        if cfg_path is not None:
            raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8")) or {}
            name = (raw.get("harness") or {}).get("timezone")
            if name:
                return ZoneInfo(str(name))
    except Exception:
        pass
    return None

_logger = logging.getLogger("openloom.storage")

_RESULT_SUFFIX_BY_EXT: dict[str, str] = {
    ".json": ".result.json",
    ".yaml": ".result.yaml",
    ".yml": ".result.yml",
    ".docx": ".result.docx",
}

_STATUS_SUFFIX_BY_EXT: dict[str, str] = {
    ".json": ".status.json",
    ".yaml": ".status.yaml",
    ".yml": ".status.yml",
    ".docx": ".status.docx",
}


def parse_spec(raw: bytes, filepath: str) -> dict[str, Any] | None:
    """Parse a task file into a normalized spec dict. Returns None if invalid."""
    ext = PurePosixPath(filepath).suffix.lower()
    try:
        if ext == ".json":
            data = json.loads(raw)
            return data if isinstance(data, dict) else None
        if ext in (".yaml", ".yml"):
            import yaml
            data = yaml.safe_load(raw)
            return data if isinstance(data, dict) else None
        if ext == ".docx":
            try:
                from docx import Document  # noqa: F401
            except ImportError:
                _logger.error(
                    "parse_spec: %s is a .docx file but python-docx is not installed. "
                    "Run: uv tool install openloom --with python-docx  or  "
                    "pip install openloom[docx]",
                    filepath,
                )
                return None
            return _parse_docx(raw)
    except Exception:
        return None
    return None


def status_suffix(source_file: str) -> str:
    """Return the status filename suffix matching the input format."""
    ext = PurePosixPath(source_file).suffix.lower()
    return _STATUS_SUFFIX_BY_EXT.get(ext, ".status.json")


def result_suffix(source_file: str) -> str:
    """Return the result filename suffix matching the input format."""
    ext = PurePosixPath(source_file).suffix.lower()
    return _RESULT_SUFFIX_BY_EXT.get(ext, ".result.json")


def render_result(payload: dict[str, Any], source_file: str) -> bytes:
    """Render a result payload in the same format as the input file."""
    ext = PurePosixPath(source_file).suffix.lower()
    if ext == ".json":
        return json.dumps(payload, indent=2, ensure_ascii=False).encode()
    if ext in (".yaml", ".yml"):
        import yaml
        return yaml.safe_dump(payload, allow_unicode=True).encode()
    if ext == ".docx":
        return _render_docx(payload, is_status=False).read()
    return json.dumps(payload, indent=2).encode()


def render_status(payload: dict[str, Any], source_file: str) -> bytes:
    """Render a status payload in the same format as the input file.

    For docx inputs this produces a human-readable status document
    (same renderer as results, with a "任务进行中" heading).
    """
    ext = PurePosixPath(source_file).suffix.lower()
    if ext == ".json":
        return json.dumps(payload, indent=2, ensure_ascii=False, separators=(",", ":")).encode()
    if ext in (".yaml", ".yml"):
        import yaml
        return yaml.safe_dump(payload, allow_unicode=True).encode()
    if ext == ".docx":
        return _render_docx(payload, is_status=True).read()
    return json.dumps(payload, indent=2).encode()


# ── DOCX spec parsing ───────────────────────────────────────────────────


def _parse_docx(raw: bytes) -> dict[str, Any] | None:
    """Parse first table with 2 columns (field | value) from a docx."""
    from docx import Document

    doc = Document(io.BytesIO(raw))
    tables = doc.tables
    if not tables:
        return None

    spec: dict[str, Any] = {}
    metadata: dict[str, Any] = {}
    for row in tables[0].rows:
        if len(row.cells) < 2:
            continue
        key = row.cells[0].text.strip().lower()
        value = row.cells[1].text.strip()
        if not key:
            continue
        if key in ("goal", "workspace", "cwd", "name", "title", "sessionid", "session_id"):
            if key == "sessionid":
                key = "sessionId"
            spec[key] = value
        else:
            metadata[key] = value

    if metadata:
        spec["metadata"] = metadata
    return spec or None


# ── DOCX rendering (human-readable) ─────────────────────────────────────


_STATUS_HEADING = {
    "completed": "任务完成报告",
    "failed": "任务失败报告",
    "running": "任务进行中",
    "waiting": "任务等待中",
    "archived": "任务已归档",
}

_TOOL_GLYPH = {
    "completed": "✓",
    "running": "▶",
    "pending": "…",
    "active": "▶",
    "failed": "✗",
    "error": "✗",
}


def _render_docx(payload: dict[str, Any], *, is_status: bool) -> io.BytesIO:
    """Render a payload as a human-readable docx.

    Structure: heading → metadata block → summary → agent trace.
    """
    from docx import Document

    doc = Document()

    status = str(payload.get("status") or "").lower()
    task_name = str(payload.get("task_name") or payload.get("task_id") or "")

    if is_status:
        heading = _STATUS_HEADING.get(status, "任务状态")
        doc.add_heading(heading, level=1)
        if task_name:
            doc.add_heading(task_name, level=2)
    else:
        heading = _STATUS_HEADING.get(status, "OpenLoom 任务结果")
        doc.add_heading(heading, level=1)
        if task_name:
            doc.add_heading(task_name, level=2)

    _docx_meta_block(doc, payload)

    data = payload.get("data") or {}
    if not isinstance(data, dict):
        data = {}

    summary = str(data.get("summary") or "").strip()
    if summary:
        doc.add_heading("概要", level=2)
        doc.add_paragraph(summary)

    recent = data.get("recent_activity")
    if isinstance(recent, list) and recent:
        _docx_trace(doc, recent)

    leftover = _docx_leftover(data)
    if leftover:
        doc.add_heading("其他信息", level=2)
        table = doc.add_table(rows=0, cols=2)
        for k, v in leftover.items():
            row = table.add_row().cells
            row[0].text = str(k)
            row[1].text = str(v)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _docx_meta_block(doc: Any, payload: dict[str, Any]) -> None:
    lines: list[str] = []
    status = str(payload.get("status") or "")
    task_id = str(payload.get("task_id") or "")
    timestamp = payload.get("timestamp")
    if status:
        lines.append(f"状态: {status}")
    if task_id:
        lines.append(f"任务 ID: {task_id}")
    if isinstance(timestamp, (int, float)) and timestamp > 0:
        iso = datetime.fromtimestamp(timestamp, tz=_resolve_tz()).strftime("%Y-%m-%d %H:%M:%S")
        lines.append(f"时间: {iso}")
    if lines:
        para = doc.add_paragraph()
        for line in lines:
            run = para.add_run(line + "\n")
            run.font.size = None
        para.runs[0].bold = True


def _docx_trace(doc: Any, recent: list[Any]) -> None:
    doc.add_heading(f"Agent 执行轨迹（共 {len(recent)} 条）", level=2)
    for idx, entry in enumerate(recent, start=1):
        if not isinstance(entry, dict):
            continue
        ts = ""
        completed = entry.get("completed_at")
        if isinstance(completed, (int, float)) and completed > 0:
            ts = datetime.fromtimestamp(
                completed / 1000 if completed > 1e12 else completed, tz=_resolve_tz(),
            ).strftime("%H:%M:%S")
        head = doc.add_paragraph()
        head_run = head.add_run(f"{idx}. {ts}".rstrip())
        head_run.bold = True
        text = str(entry.get("text") or "").strip()
        if text:
            doc.add_paragraph(text)
        tools = entry.get("tools") or []
        if isinstance(tools, list) and tools:
            for tool in tools:
                if not isinstance(tool, dict):
                    continue
                _docx_tool_line(doc, tool)


def _docx_tool_line(doc: Any, tool: dict[str, Any]) -> None:
    name = str(tool.get("tool") or tool.get("name") or "tool")
    status = str(tool.get("status") or "unknown").lower()
    glyph = _TOOL_GLYPH.get(status, "•")
    excerpt = str(tool.get("input_excerpt") or "")
    line = f"    {glyph} {name} [{status}]"
    if excerpt:
        line += f"  {excerpt}"
    para = doc.add_paragraph()
    run = para.add_run(line)
    run.font.name = "Consolas"


def _docx_leftover(data: dict[str, Any]) -> dict[str, Any]:
    used = {"summary", "recent_activity"}
    return {k: v for k, v in data.items() if k not in used}
